# the main API I am using is documented below
# https://pypi.org/project/python-docx/
# https://python-docx.readthedocs.io/en/latest/index.html

from datetime import date
from docx import Document
from docx.shared import Pt
import os

full_department_names = {
    "mech": "Department of Mechanical and Mechatronics Engineering",
    "ece": "Department of Electrical and Computer Engineering",
    "sys": "Department of Systems Design Engineering",
    "civil": "Department of Civil, Eniromental and Geological Engineering",
    "manage": "Department of Management Engineering",
    "chemical": "Department of Chemical Enineering"
}

cur_dir = os.getcwd()
waterloo_logo_file_name = "waterloo_logo.png"
path_to_waterloo_logo = os.path.join(cur_dir, "props", waterloo_logo_file_name) # where the logo lives

def format_docx_file(request):

    # read the parameters passed in
    student_name = request["studentName"]
    subject = request["subject"]
    department = full_department_names[request["department"]] # department name
    street_address = request["street"]
    postal_address = request["postal"]
    file_name = request["reportFile"]

    #############################################################
    ######################## Start job ##########################
    #############################################################
    
    path_to_file = os.path.join("workshop", "tmp", file_name) # where to read file from
    document = Document(path_to_file) # open file as a document object

    # set the deafult style (called 'Normal') to Times New Roman size 11
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    #############################################################
    ######################## COVER PAGE #########################
    #############################################################

    # start by adding the university logo
    document.add_picture(path_to_waterloo_logo)

    # then the department name
    document.add_paragraph(department + "\n\n")

    # then the title
    title = document.add_paragraph()
    title_run = title.add_run(subject + "\n\n")
    title_run.bold = True
    title_run.font.size = Pt(14)

    # then prepared for UW
    prepared_for = document.add_paragraph()
    prepared_for_run = prepared_for.add_run("A Report Prepared For:\n")
    prepared_for_run.bold = True
    prepared_for.add_run("The University of Waterloo\n")

    # then student name
    prepared_by = document.add_paragraph()
    prepared_by_run = prepared_by.add_run("Prepared By:\n")
    prepared_by_run.bold = True
    prepared_by.add_run(student_name + "\n")
    prepared_by.add_run(street_address + "\n")
    prepared_by.add_run(postal_address + "\n")

    document.add_page_break() # add page break at end of cover page

    #############################################################
    ######################## Finish and save ####################
    #############################################################
    
    final_file_name = "{} - {}.docx".format(student_name, subject)
    path_to_output = os.path.join("workshop", "output", final_file_name) # where to output file
    document.save(path_to_output) # save file in desired dest

# sample request
req = {
    "studentName": "Christian Mourad",
    "subject": "How to Format WKRPTs",
    "department": "mech",
    "street": "123 Street St",
    "postal": "City, ON P0S 1A1",
    "reportFile": "test.docx"
}

format_docx_file(req)